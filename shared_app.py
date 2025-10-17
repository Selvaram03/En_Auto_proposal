# shared_app.py
import re
from docx import Document
from io import BytesIO
import tempfile
import os
import streamlit as st

# ----------------------------
# Excel Handling
# ----------------------------
def load_excel(uploaded_file):
    """Read Excel and return dataframe with proper column checks."""
    import pandas as pd
    try:
        df = pd.read_excel(uploaded_file, engine="openpyxl")
        df.columns = df.columns.str.strip()
        if 'Parameters' not in df.columns or 'Value' not in df.columns:
            st.error("❌ Excel must have 'Parameters' and 'Value' columns.")
            return None
        df["Parameters"] = df["Parameters"].astype(str).str.strip()
        df["Value"] = df["Value"].astype(str)
        return df
    except Exception as e:
        st.error(f"❌ Error reading Excel: {e}")
        return None

# ----------------------------
# Placeholder Replacement
# ----------------------------
def replace_placeholders(doc, param_dict):
    """
    Replace placeholders {{Parameter}} in:
      - Paragraphs
      - Tables (including nested)
      - Headers and Footers
      - Text boxes (shapes)
    Handles empty keys, None values, and special characters.
    """
    pattern_template = r"\{\{\s*{}\s*\}\}"

    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            full_text = "".join(run.text for run in p.runs)
            new_text = full_text
            for key, val in param_dict.items():
                if key is None or str(key).strip() == "":
                    continue  # skip empty keys
                key_str = str(key).strip()
                val_str = "" if val is None else str(val)
                try:
                    pattern = pattern_template.format(re.escape(key_str))
                    new_text = re.sub(pattern, val_str, new_text, flags=re.IGNORECASE)
                except Exception:
                    continue
            if new_text != full_text:
                for run in p.runs:
                    run.text = ""
                p.add_run(new_text)

    # 1️⃣ Main paragraphs
    replace_in_paragraphs(doc.paragraphs)

    # 2️⃣ Tables (recursive)
    def replace_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)
                for nested_table in cell.tables:
                    replace_in_table(nested_table)
    for table in doc.tables:
        replace_in_table(table)

    # 3️⃣ Headers and Footers
    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs)
        replace_in_paragraphs(section.footer.paragraphs)
        for table in section.header.tables + section.footer.tables:
            replace_in_table(table)

    # 4️⃣ Inline shapes / text boxes (if present)
    if hasattr(doc, 'inline_shapes'):
        for shape in doc.inline_shapes:
            try:
                if shape.text_frame:
                    for paragraph in shape.text_frame.paragraphs:
                        replace_in_paragraphs([paragraph])
            except AttributeError:
                continue

# ----------------------------
# Generate Word Document
# ----------------------------
def generate_word(template_path, df):
    doc = Document(template_path)
    param_dict = {p: v for p, v in zip(df["Parameters"], df["Value"])}
    replace_placeholders(doc, param_dict)
    return doc

# ----------------------------
# Download Helpers
# ----------------------------
def download_docx(doc, filename):
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button(
        label=f"⬇️ Download {filename}.docx",
        data=buffer,
        file_name=f"{filename}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def download_pdf(doc, filename):
    try:
        from docx2pdf import convert
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, "temp.docx")
            pdf_path = os.path.join(tmpdir, "temp.pdf")
            doc.save(docx_path)
            convert(docx_path, pdf_path)
            with open(pdf_path, "rb") as f:
                st.download_button(
                    label=f"⬇️ Download {filename}.pdf",
                    data=f,
                    file_name=f"{filename}.pdf",
                    mime="application/pdf"
                )
    except Exception:
        st.warning("⚠️ PDF conversion skipped (requires MS Word or LibreOffice).")
