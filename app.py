import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from io import BytesIO
import re, tempfile, os, importlib.util, subprocess
from lxml import etree  # Needed for XML manipulation

# ========== Auto-install dependencies if missing ==========
required_libs = ["openpyxl", "lxml"]
for lib in required_libs:
    if importlib.util.find_spec(lib) is None:
        try:
            subprocess.run(["pip", "install", lib], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        except subprocess.CalledProcessError:
            st.warning(f"Could not auto-install {lib}. Please install it manually.")

# ========== Streamlit Config ==========
st.set_page_config(page_title="Proposal Auto Generator", layout="wide")
st.title("📄 Techno-Commercial Proposal Auto Generator")

# ========== SESSION STATE SETUP ==========
if 'prev_template' not in st.session_state:
    st.session_state.prev_template = None
if 'excel_uploaded' not in st.session_state:
    st.session_state.excel_uploaded = None
if 'generated_doc' not in st.session_state:
    st.session_state.generated_doc = None

# ========== Sidebar Layout ==========
enrich_logo_path = r"enrich_logo.png"

# 1️⃣ Logo first
try:
    st.sidebar.image(enrich_logo_path, width=150)
except Exception:
    st.sidebar.warning("⚠️ Logo not found or could not be loaded.")

# 2️⃣ Template selection next
st.sidebar.header("⚙️ Select Template Type")
template_choice = st.sidebar.radio(
    "Choose Template:",
    ("EPC Template", "BESS Template")
)

# ========== Reset previous data if template changed ==========
if st.session_state.prev_template != template_choice:
    st.session_state.excel_uploaded = None
    st.session_state.generated_doc = None
    st.session_state.prev_template = template_choice
    st.info("🔄 Template changed — previous uploaded data cleared. Please upload new Excel for this template.")

# ========== Set Template Paths Based on Choice ==========
if template_choice == "EPC Template":
    TEMPLATE_PATH = "EPC_Template.docx"
    TEMPLATE_EXCEL_PATH = "Input_EPC_Proposal.xlsx"
else:
    TEMPLATE_PATH = "BESS_Template.docx"
    TEMPLATE_EXCEL_PATH = "Input_BESS_Proposal.xlsx"

# ========== Active Template Banner ==========
st.markdown(
    f"""
    <div style="background-color:#f0f8ff;padding:12px;border-radius:10px;margin-bottom:10px;">
        <b>🧩 Currently Selected Template:</b> 
        <span style="color:#0056b3;">{template_choice}</span>
    </div>
    """,
    unsafe_allow_html=True
)

# ========== Instructions ==========
st.markdown("""
Upload your Excel sheet with **Parameters** and **Value** columns.  
The app will automatically replace placeholders like `{{Parameter Name}}` in the Word template,  
including those inside **text boxes**, **headers**, and **footers**.
""")

# ========== Excel Template Download ==========
try:
    with open(TEMPLATE_EXCEL_PATH, "rb") as f:
        st.download_button(
            label=f"📥 Download {template_choice} Excel Template",
            data=f,
            file_name=os.path.basename(TEMPLATE_EXCEL_PATH),
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    st.markdown("**ℹ️ Fill in this template and upload it below to generate your Word proposal.**")
except FileNotFoundError:
    st.warning(f"⚠️ {template_choice} Excel template not found at the specified path.")

# ========== File Upload ==========
uploaded_excel = st.file_uploader("📤 Upload Excel File", type=["xlsx"])
if uploaded_excel is not None:
    st.session_state.excel_uploaded = uploaded_excel

# ========== Helper Functions ==========
def replace_in_xml(doc_part, param_dict):
    try:
        if doc_part.element is None:
            return
        root = doc_part.element.getroottree()
    except AttributeError:
        return

    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main', 
                  'v': 'urn:schemas-microsoft-com:vml'}
    for key, value in param_dict.items():
        placeholder = "{{" + key + "}}"
        for elem in root.xpath('//w:t|//v:t', namespaces=namespaces):
            if elem.text and placeholder.lower() in elem.text.lower():
                pattern = r"\{\{\s*" + re.escape(key) + r"\s*\}\}"
                elem.text = re.sub(pattern, value, elem.text, flags=re.IGNORECASE)

def process_paragraphs(paragraphs, param_dict):
    def replace_placeholders(text):
        for key, value in param_dict.items():
            pattern = r"\{\{\s*" + re.escape(key) + r"\s*\}\}"
            text = re.sub(pattern, value, text, flags=re.IGNORECASE)
        return text
    for para in paragraphs:
        if "{{" in para.text:
            full_text = "".join(run.text for run in para.runs)
            new_text = replace_placeholders(full_text)
            if new_text != full_text:
                first_run = para.runs[0] if para.runs else None
                font_size = first_run.font.size if first_run and first_run.font.size else None
                for r in para.runs:
                    r.text = ""
                new_run = para.add_run(new_text)
                new_run.font.name = 'Calibri'
                if font_size:
                    new_run.font.size = font_size

def process_cell(cell, param_dict):
    process_paragraphs(cell.paragraphs, param_dict)
    for nested_table in cell.tables:
        for row in nested_table.rows:
            for cell in row.cells:
                process_cell(cell, param_dict)

def fill_template(df, template_path):
    param_dict = {p.lower(): v for p, v in zip(df["Parameters"], df["Value"])}
    doc = Document(template_path)

    replace_in_xml(doc, param_dict)
    process_paragraphs(doc.paragraphs, param_dict)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_cell(cell, param_dict)

    for section in doc.sections:
        replace_in_xml(section.header, param_dict)
        replace_in_xml(section.footer, param_dict)
        process_paragraphs(section.header.paragraphs, param_dict)
        process_paragraphs(section.footer.paragraphs, param_dict)
        try:
            for table in section.header.tables + section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        process_cell(cell, param_dict)
        except AttributeError:
            pass

    return doc

# ========== Main Logic ==========
if st.session_state.excel_uploaded is not None:
    try:
        df = pd.read_excel(st.session_state.excel_uploaded, engine="openpyxl")
        df.columns = df.columns.str.strip()
        if 'Parameters' not in df.columns or 'Value' not in df.columns:
            st.error("❌ The Excel must have 'Parameters' and 'Value' columns.")
            st.stop()

        df["Parameters"] = df["Parameters"].astype(str).str.strip()
        df["Value"] = df["Value"].astype(str)
        st.success("✅ Excel loaded successfully!")
        st.dataframe(df)

        if st.button("🚀 Generate Word Proposal"):
            try:
                filled_doc = fill_template(df, TEMPLATE_PATH)
                st.session_state.generated_doc = filled_doc  # store in session

                buffer = BytesIO()
                filled_doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label=f"⬇️ Download {template_choice.replace(' Template','')} Word File",
                    data=buffer,
                    file_name=f"Generated_{template_choice.replace(' Template','')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                try:
                    from docx2pdf import convert
                    with tempfile.TemporaryDirectory() as tmpdir:
                        docx_path = os.path.join(tmpdir, "temp.docx")
                        pdf_path = os.path.join(tmpdir, "temp.pdf")
                        filled_doc.save(docx_path)
                        convert(docx_path, pdf_path)
                        with open(pdf_path, "rb") as pdf_file:
                            st.download_button(
                                label=f"⬇️ Download {template_choice.replace(' Template','')} PDF File",
                                data=pdf_file,
                                file_name=f"Generated_{template_choice.replace(' Template','')}.pdf",
                                mime="application/pdf"
                            )
                except Exception:
                    st.warning("⚠️ PDF conversion skipped (requires MS Word or LibreOffice).")

            except Exception as e:
                st.error(f"❌ Error generating proposal: {e}")

    except Exception as e:
        st.error(f"❌ Error reading Excel: {e}. Ensure the sheet is valid and has the correct columns.")
else:
    st.info("📥 Please upload your Excel file to begin.")
